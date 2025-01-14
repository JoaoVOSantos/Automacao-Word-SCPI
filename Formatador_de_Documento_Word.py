import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
import sys
import os

from win32com.client import Dispatch  # Necessário para corrigir o arquivo Word

# Verifica se o script está rodando como um exe
if getattr(sys, 'frozen', False):
    caminho_modelo = os.path.join(sys._MEIPASS, 'Modelo.docx')
else:
    caminho_modelo = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Modelo.docx')


def corrigir_arquivo_docx(caminho):
    """
    Corrige um arquivo .docx inválido verificando sua estrutura e recriando as partes ausentes.
    """
    try:
        caminho_absoluto = os.path.abspath(caminho)
        word = Dispatch("Word.Application")
        word.Visible = False  # Deixa o processo invisível para o usuário
        doc = word.Documents.Open(caminho_absoluto)
        caminho_corrigido = caminho.replace('.docx', '_corrigido.docx')
        doc.SaveAs(caminho_corrigido, FileFormat=16)  # Salva como um DOCX válido
        doc.Close()
        word.Quit()
        print(f"Documento corrigido salvo em {caminho_corrigido}")
        return caminho_corrigido

    except Exception as e:
        print(f"Erro ao corrigir o documento com o Word: {str(e)}")
        return None
    
                
# Função para estilizar os parágrafos
def estilizar_paragrafos(doc):
    for par in doc.paragraphs:
        for run in par.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)
        par.paragraph_format.left_indent = Pt(0)
        par.paragraph_format.right_indent = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        par.paragraph_format.line_spacing = 1.15

# Função para definir margens e cabeçalhos
def definir_margens(doc):
    for section in doc.sections:
        section.top_margin = Pt(2.23 * 28.35)
        section.bottom_margin = Pt(1.5 * 28.35)
        section.left_margin = Pt(2 * 28.35)
        section.right_margin = Pt(1.5 * 28.35)
        section.header_distance = Pt(0.5 * 28.35)
        section.footer_distance = Pt(0.56 * 28.35)

# Função para formatar os nomes com a primeira letra maiúscula
def formatar_nomes(celula):
    for par in celula.paragraphs:
        # Formatar texto: primeira letra maiúscula
        par.text = " ".join([palavra.capitalize() for palavra in par.text.split()])


# Função para estilizar a Tabela 1
def estilizar_tabela1(tabela):
    # Detectar o índice da coluna "Nome"
    nome_coluna_index = None
    for idx, celula in enumerate(tabela.rows[0].cells):
        if "Nome" in celula.text:
            nome_coluna_index = idx
            break

    # Ajustar as larguras das colunas conforme especificado
    larguras = {
        0: 1.51,  # Portaria
        1: 2.00,  # Data
        2: 5.25,  # Nome
        3: 3.75,  # Cargo
        4: 2.75,  # CPF
        5: 2.26   # RG
    }

    for i, linha in enumerate(tabela.rows):
        for idx, celula in enumerate(linha.cells):
            # Definir largura das colunas
            if idx in larguras:
                largura = larguras[idx]
                celula._element.get_or_add_tcPr().append(
                    parse_xml(
                        f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{int(largura * 567)}" w:type="dxa"/>' 
                    )
                )

            # Centralizar o texto horizontalmente
            for par in celula.paragraphs:
                par.alignment = 1  # Centralizado
                for run in par.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                par.paragraph_format.line_spacing = 1.15
                par.paragraph_format.space_before = Pt(6)
            celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Verificar e formatar se for a coluna "Nome"
            if idx == nome_coluna_index:
                formatar_nomes(celula)

            # Aplicar bordas em todas as células
            tc_pr = celula._element.get_or_add_tcPr()
            borders = parse_xml(
                r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'</w:tcBorders>' 
            )
            tc_pr.append(borders)

    # Estilizar a tabela inteira após a alteração dos nomes
    for linha in tabela.rows:
        for celula in linha.cells:
            for par in celula.paragraphs:
                for run in par.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                par.alignment = 1
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                par.paragraph_format.line_spacing = 1.15
                par.paragraph_format.space_before = Pt(6)

def capitalizar_nome(nome):
    # Lista de exceções
    excecoes = ["LTDA", "EPP", "S.A.", "S.A", "SA", "ME"]
    
    # Quebra o nome em palavras e capitaliza cada uma, exceto as exceções
    palavras = nome.split()
    palavras_capitalizadas = []
    
    for palavra in palavras:
        if palavra.upper() in excecoes:
            palavras_capitalizadas.append(palavra.upper())  # Mantém a exceção em maiúsculo
        else:
            palavras_capitalizadas.append(palavra.capitalize())  # Capitaliza a palavra normalmente
    
    # Junta as palavras novamente em uma string
    return " ".join(palavras_capitalizadas)


def estilizar_tabela2(tabela):
    # Alterar nome da primeira coluna para "Lances" (apenas uma vez)
    if tabela.rows:
        cabeçalho = tabela.rows[0]
        if len(cabeçalho.cells) > 0:
            cabeçalho.cells[0].text = "Lances"
    
    # Definir as larguras das colunas em centímetros
    larguras_colunas = [1.26, 7.00, 3.00, 3.50, 2.76]  # Larguras em cm

    # Definir larguras das colunas em pontos
    larguras_colunas_pts = [largura * 28.35 for largura in larguras_colunas]

    for linha in tabela.rows:
        # Definir a largura das colunas
        for col_idx, celula in enumerate(linha.cells):
            if col_idx < len(larguras_colunas_pts):
                largura_em_pontos = int(larguras_colunas_pts[col_idx])
                celula._element.get_or_add_tcPr().append(parse_xml(
                    f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{largura_em_pontos}" w:type="dxa"/>'))

    for linha in tabela.rows:
        for idx, celula in enumerate(linha.cells):
            # Se a célula for da coluna de índice 0 (Lances), remover espaços vazios
            if idx == 0:
                for par in celula.paragraphs:
                    par.text = ' '.join(par.text.split())  # Remove espaços extras

            # Se a célula for da coluna de índice 1 (que contém os nomes), aplicar a capitalização
            if idx == 1:
                # Preservar quebras de linha
                novo_texto = []
                for par in celula.paragraphs:
                    # Capitaliza cada parágrafo, preservando a quebra de linha
                    par.text = capitalizar_nome(par.text)
                    novo_texto.append(par.text)
                # Recria o texto da célula com as quebras de linha
                celula.text = '\n'.join(novo_texto)

            for par in celula.paragraphs:
                for run in par.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                par.alignment = 1  # Centraliza o texto horizontalmente
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                par.paragraph_format.line_spacing = 1.15
                par.paragraph_format.space_before = Pt(6)

            # Adicionar bordas nas células
            tc_pr = celula._element.get_or_add_tcPr()
            borders = parse_xml(
                r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'</w:tcBorders>' 
            )
            tc_pr.append(borders)

            # Centralizar verticalmente a célula
            celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Estilizar a tabela inteira após a alteração das bordas
    for linha in tabela.rows:
        for celula in linha.cells:
            for par in celula.paragraphs:
                for run in par.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                par.alignment = 1  # Centraliza o texto horizontalmente
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                par.paragraph_format.line_spacing = 1.15
                par.paragraph_format.space_before = Pt(6)


def ajustar_largura_por_tabela(tabela):
    """
    Ajusta as larguras das colunas com base na tabela.
    A largura das colunas será configurada conforme o tipo de tabela.
    """
    # Definir larguras para a tabela com as colunas: Item, Lote, Descrição, Valor Total, Status Lance
    larguras_tabela1 = [1.21, 1.8, 10.5, 2, 2.01]

    # Verifica se a tabela tem as colunas específicas e aplica as larguras
    if len(tabela.columns) >= 5:  # Verifica se há pelo menos 5 colunas
        # Verifica os títulos das colunas para identificar a estrutura
        colunas_identificadas = [tabela.rows[0].cells[i].text.strip() for i in range(5)]
        
        # Se a estrutura for igual à esperada, define as larguras
        if colunas_identificadas == ['Item', 'Lote', 'Descrição', 'Valor Total', 'Status Lance']:
            for idx, largura in enumerate(larguras_tabela1):
                if idx < len(tabela.columns):
                    for celula in tabela.columns[idx].cells:
                        celula.width = Cm(largura)

def capitalizar_nome_com_quebra_linha(nome):
    # Lista de exceções
    excecoes = ["LTDA", "EPP", "S.A.", "S.A", "SA", "ME"]
    
    # Quebra o nome em palavras, preservando as quebras de linha
    partes = nome.split("\n")
    partes_capitalizadas = []

    for parte in partes:
        palavras = parte.split()
        palavras_capitalizadas = []

        for palavra in palavras:
            # Verifica se a palavra é uma exceção
            if palavra.upper() in excecoes:
                palavras_capitalizadas.append(palavra.upper())  # Mantém a exceção em maiúsculo
            else:
                palavras_capitalizadas.append(palavra.capitalize())  # Capitaliza a palavra normalmente

        # Junta as palavras capitalizadas novamente
        partes_capitalizadas.append(" ".join(palavras_capitalizadas))
    
    # Junta novamente as partes com quebra de linha
    return "\n".join(partes_capitalizadas)

def estilizar_tabela3(tabela):
    # Determinar as colunas a remover ou estilizar
    col_lote_idx = None
    col_codigo_idx = None
    col_valor_unitario_idx = None
    col_unidade_idx = None
    col_quantidade_idx = None
    col_marca_idx = None
    col_descricao_idx = None  # Para armazenar o índice da coluna "Descrição"

    # Verifica a primeira linha como cabeçalho para identificar as colunas
    if tabela.rows:
        for idx, celula in enumerate(tabela.rows[0].cells):
            texto_celula = celula.text.strip()
            if texto_celula.startswith("Lote"):
                col_lote_idx = idx
            elif texto_celula.startswith("Código"):
                col_codigo_idx = idx
            elif texto_celula.startswith("Valor Unitário"):
                col_valor_unitario_idx = idx
            elif texto_celula.startswith("Unidade"):
                col_unidade_idx = idx
            elif texto_celula.startswith("Quantidade"):
                col_quantidade_idx = idx
            elif texto_celula.startswith("Marca"):
                col_marca_idx = idx
            elif texto_celula.startswith("Descrição"):
                col_descricao_idx = idx  # Identificar o índice da coluna "Descrição"

    # Ajusta o índice da coluna "Lote" se estiver misturado com "Código"
    if col_codigo_idx is not None and col_lote_idx is None:
        for linha in tabela.rows:
            texto_celula = linha.cells[col_codigo_idx].text.strip()
            if "Lote" in texto_celula.split("\n")[0]:
                col_lote_idx = col_codigo_idx
                col_codigo_idx = None
                break

    # Se existir uma coluna "Código", removê-la
    if col_codigo_idx is not None:
        for linha in tabela.rows:
            linha.cells[col_codigo_idx]._element.getparent().remove(linha.cells[col_codigo_idx]._element)

    # Recalcular os índices das colunas após a remoção de "Código"
    col_lote_idx, col_valor_unitario_idx, col_unidade_idx, col_quantidade_idx, col_marca_idx, col_descricao_idx = None, None, None, None, None, None
    if tabela.rows:
        for idx, celula in enumerate(tabela.rows[0].cells):
            texto_celula = celula.text.strip()
            if texto_celula.startswith("Lote"):
                col_lote_idx = idx
            elif texto_celula.startswith("Valor Unitário"):
                col_valor_unitario_idx = idx
            elif texto_celula.startswith("Unidade"):
                col_unidade_idx = idx
            elif texto_celula.startswith("Quantidade"):
                col_quantidade_idx = idx
            elif texto_celula.startswith("Marca"):
                col_marca_idx = idx
            elif texto_celula.startswith("Descrição"):
                col_descricao_idx = idx  # Identificar novamente o índice da coluna "Descrição"

    # Ajusta larguras com base nas colunas da tabela
    ajustar_largura_por_tabela(tabela)

    # Estilizar a tabela restante
    for linha in tabela.rows:
        for idx, celula in enumerate(linha.cells):
            # Aplicar bordas em todas as células
            tc_pr = celula._element.get_or_add_tcPr()
            borders = parse_xml(
                r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'</w:tcBorders>' 
            )
            tc_pr.append(borders)

            # Centralizar verticalmente a célula
            celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Estilizar todas as células restantes
    for linha in tabela.rows:
        for idx, celula in enumerate(linha.cells):
            for par in celula.paragraphs:
                for run in par.runs:
                    run.font.name = 'Arial'  # Garantir que a fonte seja Arial
                    run.font.size = Pt(9)
                par.alignment = 1  # Centralizar texto horizontalmente
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                par.paragraph_format.line_spacing = 1.15
                par.paragraph_format.space_before = Pt(6)

            # Aplicar a função capitalizar_nome_com_quebra_linha na coluna "Descrição"
            if idx == col_descricao_idx:
                # Para cada célula na coluna "Descrição", capitalizar o texto preservando a quebra de linha
                celula.text = capitalizar_nome_com_quebra_linha(celula.text)

                # Após a capitalização, aplicar novamente a estilização na célula
                for par in celula.paragraphs:
                    for run in par.runs:
                        run.font.name = 'Arial'  # Garantir que a fonte seja Arial
                        run.font.size = Pt(10)  # Manter o tamanho da fonte
                    par.alignment = 1  # Centralizar texto horizontalmente
                    par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                    par.paragraph_format.line_spacing = 1.15
                    par.paragraph_format.space_before = Pt(6)

                
def estilizar_tabela4(tabela):
     # Para outras tabelas, você pode aplicar um estilo genérico ou personalizado
    ajustar_largura_por_tabela(tabela)  # Ajustar larguras de tabela
    for linha in tabela.rows:
        for celula in linha.cells:
             for par in celula.paragraphs:
                for run in par.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(9)
                par.alignment = 1  # Centraliza o texto horizontalmente
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                par.paragraph_format.line_spacing = 1.15
            # Adicionar bordas em todas as células
                tc_pr = celula._element.get_or_add_tcPr()
                borders = parse_xml(
                    r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                    r'</w:tcBorders>'
                )
                tc_pr.append(borders)

                # Centralizar verticalmente a célula
                celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    
def estilizar_tabela5(tabela):
    """
    Estiliza a tabela conforme as especificações:
    - Remove as colunas "Código", "Unidade" e qualquer célula contendo "Desconto".
    - Mantém a coluna "Descrição do Lote" mesmo que contenha a palavra "Código".
    - Ajusta textos das colunas "Descrição do Produto/Serviço" e "Descrição do Lote"
      para manter exceções (como LTDA, EPP, etc.) em maiúsculas e capitalizar as demais palavras.
    """
    colunas_para_remover = []
    colunas_descricao = []  # Para armazenar índices de colunas "Descrição do Produto/Serviço" e "Descrição do Lote"

    if tabela.rows:
        # Identifica as colunas para remoção e capitalização
        for idx, celula in enumerate(tabela.rows[0].cells):
            texto_celula = celula.text.strip().lower()
            if ("código" in texto_celula and "lote" not in texto_celula) or "unidade" in texto_celula or "desconto" in texto_celula:
                colunas_para_remover.append(idx)
            elif "descrição do produto/serviço" in texto_celula or "descrição do lote" in texto_celula:
                colunas_descricao.append(idx)
        
        # Ajustar texto nas colunas de descrição antes de modificar a tabela
        for linha in tabela.rows:
            for idx in colunas_descricao:
                if idx < len(linha.cells):  # Garante que o índice ainda é válido
                    celula = linha.cells[idx]
                    celula.text = capitalizar_nome(celula.text)  # Aplica a função de capitalização
                    # Reaplica a estilização para o texto ajustado
                    for par in celula.paragraphs:
                        for run in par.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(10)
                        par.alignment = 1
                        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                        par.paragraph_format.line_spacing = 1.15

        # Remove as colunas identificadas
        for idx in sorted(colunas_para_remover, reverse=True):
            for linha in tabela.rows:
                if idx < len(linha.cells):  # Verifica se o índice ainda é válido
                    linha.cells[idx]._element.getparent().remove(linha.cells[idx]._element)
    
    # Estiliza a tabela
    for linha in tabela.rows:
        for celula in linha.cells:
            # Configuração de fonte e alinhamento para todas as células
            for par in celula.paragraphs:
                for run in par.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                par.alignment = 1
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                par.paragraph_format.line_spacing = 1.15

            # Adiciona bordas às células
            tc_pr = celula._element.get_or_add_tcPr()
            borders = parse_xml(
                r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'</w:tcBorders>'
            )
            tc_pr.append(borders)

            # Centraliza verticalmente as células
            celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def capitalizar_preservando_quebra_linha(celula):
    """
    Capitaliza o texto em uma célula, preservando as quebras de linha
    e mantendo as exceções (e.g., LTDA, EPP, etc.) em maiúsculas.
    """
    excecoes = {"ltda", "epp", "s.a.", "s.a", "sa", "me"}
    
    for par in celula.paragraphs:
        for run in par.runs:
            palavras = run.text.split()  # Divide o texto do run em palavras
            palavras_capitalizadas = [
                palavra.upper() if palavra.lower() in excecoes else palavra.capitalize()
                for palavra in palavras
            ]
            # Atualiza o texto do run com as palavras capitalizadas
            run.text = " ".join(palavras_capitalizadas)

        # Reaplica a estilização no parágrafo
        par.alignment = 1
        par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        par.paragraph_format.line_spacing = 1.15

        for run in par.runs:
            run.font.name = "Arial"
            run.font.size = Pt(10)


def estilizar_tabela6(tabela):
    """
    Estiliza a tabela conforme os requisitos:
    - Fonte Arial, tamanho 10, centralizado, múltiplo 1,15.
    - Adiciona bordas pretas ao redor das células.
    - Capitaliza os textos nas colunas "Descrição do Produto/Serviço" e "Descrição do Lote",
      preservando as quebras de linha.
    - Remove colunas com "Código", exceto se associadas a "Lote".
    """
    colunas_para_remover = []
    colunas_descricao = []  # Índices das colunas "Descrição do Produto/Serviço" e "Descrição do Lote"

    if tabela.rows:
        # Identifica colunas para remoção ou estilização
        for idx, celula in enumerate(tabela.rows[0].cells):
            texto_celula = celula.text.strip().lower()
            if ("código" in texto_celula and "lote" not in texto_celula) or "unidade" in texto_celula or "desconto" in texto_celula:
                colunas_para_remover.append(idx)
            elif "descrição do produto/serviço" in texto_celula or "descrição do lote" in texto_celula:
                colunas_descricao.append(idx)

        # Ajustar texto nas colunas de descrição antes de modificar a tabela
        for linha in tabela.rows:
            for idx in colunas_descricao:
                if idx < len(linha.cells):  # Verifica se o índice é válido
                    celula = linha.cells[idx]
                    capitalizar_preservando_quebra_linha(celula)  # Aplica a nova função de capitalização

        # Remove as colunas identificadas
        for idx in sorted(colunas_para_remover, reverse=True):
            for linha in tabela.rows:
                if idx < len(linha.cells):  # Garante que o índice ainda é válido
                    linha.cells[idx]._element.getparent().remove(linha.cells[idx]._element)

    # Estilização geral da tabela
    for linha in tabela.rows:
        for celula in linha.cells:
            # Configuração de fonte, alinhamento e espaçamento
            for par in celula.paragraphs:
                for run in par.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                par.alignment = 1
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                par.paragraph_format.line_spacing = 1.15

            # Adiciona bordas às células
            tc_pr = celula._element.get_or_add_tcPr()
            borders = parse_xml(
                r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
                r'</w:tcBorders>'
            )
            tc_pr.append(borders)

            # Centraliza verticalmente as células
            celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def capitalizar_inabilitado(texto):
    # Se o texto começa com "Inabilitado", capitaliza corretamente a primeira palavra
    if texto.startswith("Inabilitado"):
        partes = texto.split("\n", 1)  # Divida em Inabilitado e o resto
        partes[1] = partes[1].lower().capitalize()  # Capitaliza o resto do texto após a quebra de linha
        return partes[0] + "\n" + partes[1]
    return texto


def estilizar_tabela7(tabela):
    """
    Estiliza a tabela, aplicando as seguintes regras:
    - Na coluna "Proponente / Fornecedor", capitaliza as exceções (LTDA, ME, EPP, etc.).
    - Na coluna "Representante", capitaliza apenas a primeira letra de cada palavra.
    - Na última coluna (index 4), capitaliza a palavra 'Inabilitado' corretamente e ajusta o texto subsequente.
    - Aplica bordas, centraliza o texto e configura a fonte para Arial, tamanho 10, com espaçamento de 1,15.
    """
    for linha in tabela.rows:
        for idx, celula in enumerate(linha.cells):
            texto_celula = celula.text.strip().lower()

            # Capitaliza a coluna "Proponente / Fornecedor" usando a função capitalizar_nome
            if idx == 1:  # Proponente / Fornecedor
                celula.text = capitalizar_nome(celula.text)

            # Capitaliza a coluna "Representante" usando a função capitalizar_nome
            elif idx == 3:  # Representante
                celula.text = capitalizar_nome(celula.text)

            # Ajusta o texto da última coluna (index 4) para capitalizar corretamente 'Inabilitado'
            elif idx == 4:  # Última coluna
                celula.text = capitalizar_inabilitado(celula.text)

            # Configuração de fonte, alinhamento e espaçamento
            for par in celula.paragraphs:
                for run in par.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                par.alignment = 1
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                par.paragraph_format.line_spacing = 1.15

            # Adiciona bordas às células
            tc_pr = celula._element.get_or_add_tcPr()
            borders = parse_xml(
                r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'</w:tcBorders>' 
            )
            tc_pr.append(borders)

            # Centraliza verticalmente as células
            celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def capitalizar_nome_com_quebra(texto):
    """
    Capitaliza o nome, preservando as quebras de linha. As exceções, como LTDA, ME, etc., são mantidas em maiúsculas.
    """
    linhas = texto.split('\n')  # Divide o texto nas quebras de linha
    linhas_capitalizadas = []

    for linha in linhas:
        palavras = linha.split()  # Divide a linha em palavras
        palavras_capitalizadas = []
        
        for palavra in palavras:
            # Verifica se a palavra é uma exceção e a mantém em maiúsculas, caso contrário capitaliza a palavra
            if palavra.upper() in ["LTDA", "ME", "MEI", "EIRELI"]:  # Liste outras exceções aqui
                palavras_capitalizadas.append(palavra.upper())
            else:
                palavras_capitalizadas.append(palavra.capitalize())  # Capitaliza a palavra

        linhas_capitalizadas.append(' '.join(palavras_capitalizadas))  # Junta as palavras capitalizadas

    return '\n'.join(linhas_capitalizadas)  # Junta novamente com quebras de linha


def estilizar_tabela8(tabela):
    """
    Estiliza a tabela 8, aplicando as seguintes regras:
    - Se a primeira palavra da célula for 'Código', remove a coluna inteira.
    - Se for 'Lote', o valor será mantido.
    - Na coluna 'Descrição', capitaliza as exceções (LTDA, ME, etc.) e coloca a primeira letra de cada palavra em maiúscula, preservando as quebras de linha.
    """
    
    # Verifica e remove a coluna se a primeira célula de qualquer linha contiver 'Código'
    col_codigo = None  # Variável para armazenar o índice da coluna 'Código'
    
    for linha in tabela.rows:
        for idx, celula in enumerate(linha.cells):
            texto_celula = celula.text.strip().lower()

            # Se a primeira palavra for "Código", armazena o índice da coluna para remoção
            if texto_celula.startswith("código"):
                col_codigo = idx
                break  # Sai do loop após encontrar a coluna

    if col_codigo is not None:
        # Remove a coluna inteira se encontrada
        for i in range(len(tabela.rows)):  # Para cada linha, remove a célula na coluna
            tabela.rows[i]._element.remove(tabela.rows[i].cells[col_codigo]._tc)

    # Processa as colunas restantes
    for linha in tabela.rows:
        for idx, celula in enumerate(linha.cells):
            texto_celula = celula.text.strip().lower()

            # Capitaliza a coluna "Descrição" com exceções, preservando quebras de linha
            if idx == 2:  # Descrição
                celula.text = capitalizar_nome_com_quebra(celula.text)

            # Configuração de fonte, alinhamento e espaçamento
            for par in celula.paragraphs:
                for run in par.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10)
                par.alignment = 1
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                par.paragraph_format.line_spacing = 1.15

            # Adiciona bordas às células, igual as outras tabelas
            tc_pr = celula._element.get_or_add_tcPr()
            borders = parse_xml(
                r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'</w:tcBorders>' 
            )
            tc_pr.append(borders)

            # Centraliza verticalmente as células
            celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def estilizar_tabela9(tabela):
    """
    Estiliza a tabela 9 aplicando as seguintes regras:
    - Remove a coluna de índice 1 somente se for 'Código' (não 'Lote').
    - Aplica estilização com Arial, tamanho 9, centralizado, espaçamento múltiplo de 1.15.
    - Define larguras manuais específicas caso a coluna "Código" seja removida.
    """

    colunas_para_remover = []  # Índices das colunas a serem removidas

    # Verifica se o texto da coluna contém 'Código' ou 'Lote'
    if len(tabela.rows) > 0:  # Verifica se há linhas na tabela
        texto_primeira_coluna = tabela.rows[0].cells[1].text.strip().lower()
        texto_ultima_coluna = tabela.rows[-1].cells[1].text.strip().lower()

        # A lógica aqui verifica se a palavra "lote" não está presente em nenhuma das colunas
        if "lote" not in texto_primeira_coluna and "lote" not in texto_ultima_coluna:
            colunas_para_remover.append(1)

    # Recria a tabela sem as colunas indesejadas
    if colunas_para_remover:
        for linha in tabela.rows:
            nova_linha = [celula.text for idx, celula in enumerate(linha.cells) if idx not in colunas_para_remover]
            while len(linha.cells) > len(nova_linha):  # Remove colunas adicionais
                linha.cells[-1]._element.getparent().remove(linha.cells[-1]._element)

            # Atualiza as células restantes na linha
            for idx, texto in enumerate(nova_linha):
                linha.cells[idx].text = texto

        # Define larguras manuais para as colunas restantes
        larguras = [Cm(1.01), Cm(9.26), Cm(1.41), Cm(1.85), Cm(2.05), Cm(1.95)]
        for linha in tabela.rows:
            for idx, celula in enumerate(linha.cells):
                if idx < len(larguras):  # Aplica as larguras definidas, se disponíveis
                    celula.width = larguras[idx]

    # Aplica estilização nas colunas restantes
    for linha in tabela.rows:
        for celula in linha.cells:
            # Configuração de fonte, alinhamento e espaçamento
            for par in celula.paragraphs:
                for run in par.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(9)
                par.alignment = 1  # Alinhamento centrado horizontal
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                par.paragraph_format.line_spacing = 1.15

            # Adiciona bordas às células
            tc_pr = celula._element.get_or_add_tcPr()
            borders = parse_xml(
                r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>' 
                r'</w:tcBorders>' 
            )
            tc_pr.append(borders)

            # Centraliza verticalmente as células
            celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER



def estilizar_tabelas(doc):
    for i, tabela in enumerate(doc.tables):
        try:
            print(f"Iniciando a estilização da Tabela {i + 1}...")

            if i == 0:
                # Estilizar a primeira tabela com estilo de Tabela 1
                estilizar_tabela1(tabela)
            elif i == 1:
                # Estilizar a segunda tabela com estilo de Tabela 2
                estilizar_tabela2(tabela)
            elif i == 2:
                # Estilizar a terceira tabela com estilo de Tabela 3
                estilizar_tabela3(tabela)
            elif i == 3:
                estilizar_tabela4(tabela)
            elif i == 4:
                estilizar_tabela5(tabela)
            elif i == 5:
                estilizar_tabela6(tabela)
            elif i == 6:
                estilizar_tabela7(tabela)
            elif i == 7:
                estilizar_tabela8(tabela)
            elif i == 8:
                estilizar_tabela9(tabela)
            else:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        # Configuração de fonte, alinhamento e espaçamento
                        for par in celula.paragraphs:
                            for run in par.runs:
                                run.font.name = "Arial"
                                run.font.size = Pt(9)
                            par.alignment = 1  # Alinhamento centrado horizontal
                            par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                            par.paragraph_format.line_spacing = 1.15
                            par.paragraph_format.space_before = Pt(6)  

                        # Centraliza verticalmente as células
                        celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            print(f"Estilização da Tabela {i + 1} concluída com sucesso!")
        
        except Exception as e:
            print(f"Erro ao estilizar a Tabela {i + 1}: {str(e)}")



def excluir_arquivo(caminho):
    try:
        os.remove(caminho)
        print(f"Arquivo {caminho} excluído com sucesso.")
    except Exception as e:
        print(f"Erro ao excluir o arquivo {caminho}: {str(e)}")



# Função para copiar conteúdo e estilizar
def copiar_conteudo_para_modelo():
    caminho_arquivo = filedialog.askopenfilename(title="Selecione um arquivo Word", filetypes=[("Documentos Word", "*.docx")])
    if caminho_arquivo:
        # Corrigir e salvar o arquivo selecionado como um novo .docx
        caminho_corrigido = corrigir_arquivo_docx(caminho_arquivo)
        if not caminho_corrigido:
            messagebox.showerror("Erro", "Não foi possível corrigir o arquivo selecionado.")
            return

        doc_origem = Document(caminho_corrigido)
        
        if os.path.exists(caminho_modelo):
            doc_modelo = Document(caminho_modelo)
            for elemento in doc_origem.element.body:
                doc_modelo.element.body.append(elemento)
            
            estilizar_paragrafos(doc_modelo)
            definir_margens(doc_modelo)
            estilizar_tabelas(doc_modelo)

            caminho_saida = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Documentos Word", "*.docx")])
            if caminho_saida:
                doc_modelo.save(caminho_saida)
                print(f"Arquivo salvo em: {caminho_saida}")

                # Excluir o arquivo corrigido após salvar o novo documento
                excluir_arquivo(caminho_corrigido)
        else:
            print(f"Modelo não encontrado: {caminho_modelo}")




# Interface gráfica
def criar_interface():
    root = tk.Tk()
    root.withdraw()
    copiar_conteudo_para_modelo()

# Executa o programa
criar_interface()
