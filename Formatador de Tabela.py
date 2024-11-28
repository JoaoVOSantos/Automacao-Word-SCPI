import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn

# Função para verificar se a célula pertence à coluna "Nome"
def is_coluna_nome(tabela, index_celula):
    for linha in tabela.rows:
        if linha.cells[index_celula].text.strip().lower() == "nome":
            return True
    return False

# Função para formatar os nomes com a primeira letra maiúscula
def formatar_nomes(celula):
    for par in celula.paragraphs:
        par.text = " ".join([palavra.capitalize() for palavra in par.text.split()])

# Função para estilizar as tabelas de maneira geral
def estilizar_tabela(tabela):
    # Detectar o índice da coluna "Nome"
    nome_coluna_index = None
    for idx, celula in enumerate(tabela.rows[0].cells):
        if "Nome" in celula.text:
            nome_coluna_index = idx
            break
    
    # Ajustar as larguras das colunas conforme especificado
    larguras = {
        0: 1.51,  # Portaria
        1: 2.00,  # Data
        2: 5.25,   # Nome
        3: 3.75,  # Cargo
        4: 2.75,   # CPF
        5: 2.26   # RG
    }
    
    for i, linha in enumerate(tabela.rows):
        for idx, celula in enumerate(linha.cells):
            # Definir largura das colunas
            if idx in larguras:
                largura = larguras[idx]
                celula._element.get_or_add_tcPr().append(
                    parse_xml(
                        f'<w:tcW xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:w="{int(largura * 567)}" w:type="dxa"/>'
                    )
                )

            # Centralizar o texto horizontalmente
            for par in celula.paragraphs:
                par.alignment = 1  # 1 = centralizado
                for run in par.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                par.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                par.paragraph_format.line_spacing = 1.5
                par.paragraph_format.space_before = Pt(6)
            celula.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Verificar e formatar se for a coluna "Nome"
            if idx == nome_coluna_index:
                formatar_nomes(celula)

            # Aplicar bordas em todas as células
            for borda in celula._element.xpath('.//w:tcPr//w:tcBorders'):
                borda.clear()  # Limpar bordas existentes
            celula._element.get_or_add_tcPr().append(
                parse_xml(
                    r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                    r'<w:top w:val="single" w:space="0" w:size="4" w:color="000000"/>'
                    r'<w:left w:val="single" w:space="0" w:size="4" w:color="000000"/>'
                    r'<w:bottom w:val="single" w:space="0" w:size="4" w:color="000000"/>'
                    r'<w:right w:val="single" w:space="0" w:size="4" w:color="000000"/>'
                    r'</w:tcBorders>'
                )
            )

# Função para aplicar estilo às tabelas do documento
def estilizar_tabelas(doc):
    for i, tabela in enumerate(doc.tables):
        if i < 2:  # Estilizar as duas primeiras tabelas
            estilizar_tabela(tabela)
    return doc

# Função para abrir o arquivo Word
def abrir_arquivo():
    caminho_arquivo = filedialog.askopenfilename(
        title="Selecione um arquivo Word", filetypes=[("Documentos Word", "*.docx")]
    )
    if caminho_arquivo:
        doc = Document(caminho_arquivo)
        doc = estilizar_tabelas(doc)
        # Salvar o arquivo estilizado com novo nome
        caminho_saida = filedialog.asksaveasfilename(
            defaultextension=".docx", filetypes=[("Documentos Word", "*.docx")]
        )
        if caminho_saida:
            doc.save(caminho_saida)
            print(f"Arquivo salvo em: {caminho_saida}")
        else:
            print("Erro ao salvar o arquivo.")

# Criando a interface gráfica
root = tk.Tk()
root.withdraw()  # Esconde a janela principal

# Abre a janela para selecionar o arquivo
abrir_arquivo()
