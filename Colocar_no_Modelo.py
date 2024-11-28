import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Pt  # Importando Pt para definir o tamanho da fonte
import os

# Função para estilizar todos os parágrafos do documento
def estilizar_paragrafos(doc):
    for par in doc.paragraphs:
        # Definindo a fonte para Arial e tamanho 10
        for run in par.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(10)
        
        # Definindo o recuo para 0,0 e o espaçamento para 0,0
        par.paragraph_format.left_indent = Pt(0)
        par.paragraph_format.right_indent = Pt(0)
        par.paragraph_format.space_before = Pt(0)
        par.paragraph_format.space_after = Pt(0)

        # Definindo o espaçamento entre linhas para múltiplos de 1,15
        par.paragraph_format.line_spacing_rule = 2  # Multiplicativo
        par.paragraph_format.line_spacing = 1.15

# Função para definir as margens do documento e do cabeçalho/rodapé
def definir_margens(doc):
    sections = doc.sections
    for section in sections:
        # Definindo as margens (em pontos, 1 cm = 28,35 pontos)
        section.top_margin = Pt(2.23 * 28.35)  # Superior: 2,23cm
        section.bottom_margin = Pt(1.5 * 28.35)  # Inferior: 1,5cm
        section.left_margin = Pt(2 * 28.35)  # Esquerda: 2cm
        section.right_margin = Pt(1.5 * 28.35)  # Direita: 1,5cm
        
        # Definindo a posição do cabeçalho e rodapé
        section.header_distance = Pt(0.5 * 28.35)  # Cabeçalho: 0,5cm
        section.footer_distance = Pt(0.56 * 28.35)  # Rodapé: 0,56cm

        # Configurando a altura do cabeçalho e rodapé
        section.header_height = Pt(0.5 * 28.35)  # Cabeçalho: 0,5cm
        section.footer_height = Pt(0.56 * 28.35)  # Rodapé: 0,56cm

# Função para copiar o conteúdo para o modelo e estilizar
def copiar_conteudo_para_modelo():
    # Seleciona o documento de origem
    caminho_arquivo = filedialog.askopenfilename(
        title="Selecione um arquivo Word", filetypes=[("Documentos Word", "*.docx")]
    )
    
    if caminho_arquivo:
        # Abre o documento de origem
        doc_origem = Document(caminho_arquivo)

        # Define o caminho do modelo (Modelo.docx) automaticamente a partir do diretório do script
        caminho_modelo = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Modelo.docx')
        
        if os.path.exists(caminho_modelo):
            # Abre o modelo
            doc_modelo = Document(caminho_modelo)
            
            # Copiar o conteúdo do documento de origem para o modelo
            for elemento in doc_origem.element.body:
                doc_modelo.element.body.append(elemento)

            # Estilizar o novo documento
            estilizar_paragrafos(doc_modelo)
            definir_margens(doc_modelo)
            
            # Salvar o novo documento
            caminho_saida = filedialog.asksaveasfilename(
                defaultextension=".docx", filetypes=[("Documentos Word", "*.docx")]
            )
            if caminho_saida:
                doc_modelo.save(caminho_saida)
                print(f"Arquivo salvo em: {caminho_saida}")
            else:
                print("Erro ao salvar o arquivo.")
        else:
            print(f"Modelo não encontrado no caminho: {caminho_modelo}")
    else:
        print("Erro ao selecionar o documento de origem.")

# Função para criar a interface gráfica
def criar_interface():
    # Cria a janela principal
    root = tk.Tk()
    root.withdraw()  # Esconde a janela principal

    # Abre a janela para selecionar os arquivos
    copiar_conteudo_para_modelo()

# Executa o programa
criar_interface()
